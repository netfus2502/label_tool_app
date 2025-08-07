import streamlit as st
import pandas as pd
import qrcode
import barcode
from barcode.writer import ImageWriter
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import Inches
import os

# 🔐 認証処理（最初に実行）
PASSWORD = "fukuoka2025"
st.title("ラベル生成ツール")

password_input = st.text_input("パスワードを入力してください", type="password")
if password_input != PASSWORD:
    st.warning("正しいパスワードを入力してください")
    st.stop()  # 認証失敗 → ここで処理を止める

# ✅ 認証成功後の処理（ここからが本体）
st.success("認証成功！ツールを利用できます")

st.set_page_config(page_title="Wordラベル生成ツール", layout="wide")
st.title("🏷️ Wordラベル生成ツール（2列×N行）")

# CSVアップロード
uploaded_file = st.file_uploader("CSVファイルをアップロードしてください", type=["csv"])
if uploaded_file:
    df = pd.read_csv(uploaded_file)
    st.success("CSVファイルを読み込みました ✅")
    st.dataframe(df)

    # 商品URL列があるか確認
    has_url = '商品URL' in df.columns
    if not has_url:
        st.warning("⚠️ CSVに『商品URL』列がありません。QRコードとURL表示は省略されます。")

    # バーコード生成
    def generate_barcode(code):
        buffer = BytesIO()
        ean = barcode.get('code128', str(code), writer=ImageWriter())
        ean.write(buffer)
        buffer.seek(0)
        return buffer

    # QRコード生成
    def generate_qr(url):
        qr = qrcode.make(url)
        buffer = BytesIO()
        qr.save(buffer, format="PNG")
        buffer.seek(0)
        return buffer

    # Wordラベル生成（2列×N行）
    def generate_word_labels(dataframe):
        doc = Document()
        doc.add_heading("商品ラベル一覧", level=1)

        table = doc.add_table(rows=0, cols=2)
        table.autofit = False

        for index, row in dataframe.iterrows():
            # 商品情報取得
            name = row.get('商品名', '')
            code = row.get('商品コード', '')
            url = row.get('商品URL', '') if has_url else ''

            # 画像生成
            barcode_buffer = generate_barcode(code)
            barcode_img = Image.open(barcode_buffer)
            barcode_path = f"{code}_barcode.png"
            barcode_img.save(barcode_path)

            # QRコード（URLがある場合のみ）
            if url:
                qr_buffer = generate_qr(url)
                qr_img = Image.open(qr_buffer)
                qr_path = f"{code}_qr.png"
                qr_img.save(qr_path)

            # ラベル内容
            label_text = f"商品名: {name}\n商品コード: {code}"
            if url:
                label_text += f"\nURL: {url}"

            # 1行に2ラベル（左と右）
            if index % 2 == 0:
                row_cells = table.add_row().cells

            cell = row_cells[index % 2]
            cell.paragraphs[0].add_run(label_text + "\n")
            cell.paragraphs[0].add_run().add_picture(barcode_path, width=Inches(2.0))
            if url:
                cell.paragraphs[0].add_run().add_picture(qr_path, width=Inches(1.2))

        # 保存
        word_buffer = BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)
        return word_buffer

    # ダウンロードボタン
    st.subheader("📄 Wordラベル出力")
    if st.button("Wordラベルを生成"):
        word_buffer = generate_word_labels(df)
        st.download_button(
            label="📥 Wordファイルをダウンロード",
            data=word_buffer,
            file_name="商品ラベル.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
